"""Convert docs/component-rationale.md to a publication-ready Word document."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

FONT = "Times New Roman"
FONT_SIZE = Pt(10)
HEADER_BG = "D9D9D9"
BLOCKQUOTE_BG = "F5F5F5"

REPO = Path("/Users/pauloldham/Documents/cali-allocation-model-unscale-v3")


def _set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_styled_run(paragraph, text, bold=False, italic=False, font_name=FONT, font_size=FONT_SIZE, color=None):
    """Add a run with specific formatting, handling inline bold/italic markdown."""
    # Parse inline markdown: **bold** and *italic*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = font_size
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color
    return paragraph


def _parse_table(lines):
    """Parse markdown table lines into list of rows (list of list of str)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator lines
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows


def _add_table(doc, rows):
    """Add a styled Word table from parsed rows."""
    if not rows:
        return
    n_cols = len(rows[0])
    n_rows = len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        is_header = (i == 0)
        # Detect bold rows (contain ** markers)
        has_bold_marker = any('**' in c for c in row_data)
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            # Clean cell text
            clean = cell_text.replace('**', '')
            run = p.add_run(clean)
            run.font.name = FONT
            run.font.size = FONT_SIZE
            if is_header:
                run.font.bold = True
                _set_cell_shading(cell, HEADER_BG)
            elif has_bold_marker:
                run.font.bold = True

    # Auto-size
    text_width = Cm(4.0)
    num_width = Cm(3.0)
    for row in table.rows:
        for j in range(n_cols):
            row.cells[j].width = text_width if j == 0 else num_width


def convert_md_to_word(md_path: Path, output_path: Path):
    """Convert the component-rationale.md to a Word document."""
    text = md_path.read_text(encoding='utf-8')

    # Skip YAML frontmatter
    if text.startswith('---'):
        end = text.find('---', 3)
        if end != -1:
            text = text[end + 3:].lstrip()

    lines = text.split('\n')

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    table_buffer = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Flush table buffer when leaving table context
        if in_table and not stripped.startswith('|') and stripped != '':
            rows = _parse_table(table_buffer)
            _add_table(doc, rows)
            table_buffer = []
            in_table = False

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            # Main title - skip (YAML already removed, this is the doc title)
            h = doc.add_heading(stripped.lstrip('# ').strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = FONT
                run.font.size = Pt(16)
            i += 1
            continue

        if stripped.startswith('## '):
            level = len(re.match(r'^#+', stripped).group())
            heading_text = stripped.lstrip('# ').strip()
            h = doc.add_heading(heading_text, level=min(level, 4))
            for run in h.runs:
                run.font.name = FONT
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add left border via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="4" w:space="4" w:color="999999"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            _add_styled_run(p, quote_text, italic=True)
            i += 1
            continue

        # Table lines
        if stripped.startswith('|'):
            table_buffer.append(stripped)
            in_table = True
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(8)
            i += 1
            continue

        # Blank line
        if stripped == '':
            i += 1
            continue

        # Formula line (contains × or mathematical notation)
        if '×' in stripped or '= (' in stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_styled_run(p, stripped, font_size=Pt(11))
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            _add_styled_run(p, m.group(2))
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_styled_run(p, stripped[2:])
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_styled_run(p, stripped)
        i += 1

    # Flush any remaining table
    if table_buffer:
        rows = _parse_table(table_buffer)
        _add_table(doc, rows)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    md_path = REPO / "docs" / "component-rationale.md"
    output_path = REPO / "model-tables" / "component-rationale.docx"
    result = convert_md_to_word(md_path, output_path)
    print(f"Created: {result}")
