"""Reusable CSV/Excel → Word table converter.

Produces a clean black-and-white Word table with:
  - Times New Roman 10pt
  - Grey header row
  - Bold Total row (auto-detected)
  - Auto-number formatting for numeric columns
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

FONT = "Times New Roman"
FONT_SIZE = Pt(10)
HEADER_BG = "D9D9D9"

# Column rename rules applied when the source column matches exactly
DEFAULT_COL_MAP = {
    "region": "UN Region",
    "intermediate_region": "Intermediate Region",
    "sub_region": "Sub-Region",
    "Countries (number)": "Countries",
    "total_allocation": "Total Allocation\n(USD M)",
    "state_component": "State Component\n(USD M)",
    "iplc_component": "IPLC Component\n(USD M)",
}


def _set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _format_val(val) -> str:
    if isinstance(val, float):
        return f"{val:,.2f}"
    if isinstance(val, int):
        return str(val)
    return str(val)


def csv_to_word(
    input_path: str | Path,
    output_path: str | Path | None = None,
    title: str | None = None,
    subtitle: str | None = None,
    col_map: dict | None = None,
    total_keyword: str = "total",
) -> Path:
    """Convert a CSV or Excel file to a styled Word document with a table.

    Parameters
    ----------
    input_path : path to .csv, .xlsx, or .xls file
    output_path : path for the .docx output (default: same stem, .docx suffix)
    title : heading above the table (default: derived from filename stem)
    subtitle : smaller italic text below the title (optional)
    col_map : dict mapping source column names to display names (default: DEFAULT_COL_MAP)
    total_keyword : string that identifies the summary/total row (case-insensitive)

    Returns
    -------
    Path to the generated .docx file
    """
    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path.with_suffix(".docx")
    output_path = Path(output_path)

    suffix = input_path.suffix.lower()
    if suffix == ".csv":
        df = pd.read_csv(input_path)
    elif suffix in (".xlsx", ".xls"):
        df = pd.read_excel(input_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    # Rename columns
    cmap = col_map if col_map is not None else DEFAULT_COL_MAP
    df = df.rename(columns={c: v for c, v in cmap.items() if c in df.columns})

    # Round float columns to 2 dp
    for c in df.columns:
        if df[c].dtype in ("float64", "float32"):
            df[c] = df[c].round(2)

    # Derive title from filename if not provided
    if title is None:
        title = input_path.stem.replace("-", " ").replace("_", " ").title()

    # --- Build document ---
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    h = doc.add_heading(title, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = FONT
        run.font.size = Pt(13)

    # Subtitle
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(subtitle)
        run.font.name = FONT
        run.font.size = FONT_SIZE
        run.font.italic = True

    doc.add_paragraph()  # spacer

    # --- Table ---
    n_rows = len(df)
    n_cols = len(df.columns)
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for j, col_name in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = str(col_name).split("\n")
        for k, line in enumerate(lines):
            if k > 0:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.name = FONT
            run.font.size = FONT_SIZE
            run.font.bold = True
        _set_cell_shading(cell, HEADER_BG)

    # Data rows
    first_col = df.columns[0]
    for i, row in df.iterrows():
        is_total = str(row[first_col]).strip().lower() == total_keyword.lower()
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_format_val(val))
            run.font.name = FONT
            run.font.size = FONT_SIZE
            if is_total:
                run.font.bold = True

    # Auto-size columns: wider for text, narrower for numbers
    text_width = Cm(4.0)
    num_width = Cm(3.0)
    for row in table.rows:
        for j in range(n_cols):
            row.cells[j].width = text_width if j == 0 else num_width

    doc.save(output_path)
    return output_path
