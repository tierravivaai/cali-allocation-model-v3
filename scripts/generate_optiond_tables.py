"""Generate Option D revision tables (band-order preservation and break-point summary).

Outputs two Word documents in model-tables/:
1. iusaf-band-order-preservation.docx — Table showing Band 6 vs Band 5 mean at each TSAC level
2. iusaf-breakpoint-summary.docx — Table showing TSAC/SOSAC/IUSAF%/Spearman/description
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from cali_model.data_loader import load_data, get_base_data
from cali_model.calculator import calculate_allocations

FONT = "Times New Roman"
HEADER_BG = "D9D9D9"
OUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'model-tables')

FUND = 1_000_000_000
IPLC = 50


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def compute_scenario(con, beta, gamma):
    """Compute allocations and metrics for a single scenario."""
    base_df = get_base_data(con)
    pure = calculate_allocations(base_df, FUND, IPLC, exclude_high_income=True,
                                 tsac_beta=0, sosac_gamma=0, equality_mode=False,
                                 un_scale_mode="band_inversion")
    pure_eligible = pure[pure['eligible']]

    df = calculate_allocations(base_df, FUND, IPLC, exclude_high_income=True,
                               tsac_beta=beta, sosac_gamma=gamma, equality_mode=False,
                               un_scale_mode="band_inversion")
    eligible = df[df['eligible']]

    # Band means
    band6 = eligible[eligible['un_band'].str.startswith('Band 6')]
    band5 = eligible[eligible['un_band'].str.startswith('Band 5')]
    b6_mean = band6['total_allocation'].mean()
    b5_mean = band5['total_allocation'].mean()
    margin = (b5_mean - b6_mean) / b5_mean * 100 if b5_mean > 0 else 0
    order_ok = b5_mean > b6_mean

    # Spearman vs pure IUSAF
    merged = eligible[['party', 'final_share']].merge(
        pure_eligible[['party', 'final_share']], on='party', suffixes=('_cur', '_base'))
    r_cur = merged['final_share_cur'].rank(method='average')
    r_base = merged['final_share_base'].rank(method='average')
    spearman = float(r_cur.corr(r_base, method='pearson'))

    iusaf_pct = (1 - beta - gamma) * 100

    return {
        'b6_mean': b6_mean, 'b5_mean': b5_mean, 'margin': margin,
        'order_ok': order_ok, 'spearman': spearman, 'iusaf_pct': iusaf_pct,
    }


def write_band_order_table(con, path):
    """Generate the band-order preservation table."""
    scenarios = [
        ('0% (Pure IUSAF)', 0.0, 0.0),
        ('1.5% (Strict)', 0.015, 0.03),
        ('2.5% (Gini-minimum)', 0.025, 0.03),
        ('3.0% (Band-order overturn)', 0.03, 0.03),
        ('3.5% (Bounded)', 0.035, 0.03),
        ('9.2% (TSAC component overturn)', 0.092, 0.03),
    ]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title = doc.add_heading("Band-Order Preservation Across TSAC Levels", level=2)
    for run in title.runs:
        run.font.name = FONT
        run.font.size = Pt(11)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("142 eligible Parties | Fund size: USD 1,000 M | SOSAC = 3% | State/IPLC split: 50/50")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.italic = True

    doc.add_paragraph()

    headers = ['TSAC Level', 'Band 6 mean\n(China)', 'Band 5 mean\n(Brazil, India, Mexico)', 'Band 5 vs\nBand 6 margin', 'IUSAF Band\nOrder Preserved?']
    n_cols = len(headers)
    n_rows = len(scenarios)
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = hdr.split("\n")
        for k, line in enumerate(lines):
            if k > 0:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.name = FONT
            run.font.size = Pt(8.5)
            run.font.bold = True
        set_cell_shading(cell, HEADER_BG)

    # Data
    for i, (label, beta, gamma) in enumerate(scenarios):
        m = compute_scenario(con, beta, gamma)
        row = table.rows[i + 1]

        values = [
            label,
            f"USD {m['b6_mean']:.2f}M",
            f"USD {m['b5_mean']:.2f}M",
            f"{m['margin']:+.1f}%",
            "YES" if m['order_ok'] and m['margin'] > 5 else
                f"YES (margin {m['margin']:.1f}%)" if m['order_ok'] else
                "NO — Band 6 overtakes Band 5",
        ]

        for j, val in enumerate(values):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(8.5)

            # Highlight Gini-minimum row green, overturn row red
            if 'Gini-minimum' in label:
                run.font.bold = True
                set_cell_shading(cell, "D4EDDA")
            elif 'overturn' in label and not m['order_ok']:
                if j == len(values) - 1:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    run.font.bold = True

    # Column widths
    widths = [Cm(5.5), Cm(3.0), Cm(4.0), Cm(2.5), Cm(4.5)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w

    doc.save(path)
    print(f"Saved: {path}")


def write_breakpoint_summary_table(con, path):
    """Generate the break-point summary table."""
    scenarios = [
        ('0%', '3%', 0.0, 0.03, 'SOSAC only — modest rank shift among SIDS'),
        ('1.5%', '3%', 0.015, 0.03, 'Strict — IUSAF dominant for all Parties'),
        ('2.5%', '3%', 0.025, 0.03, 'Gini-minimum — band order preserved (margin 5.4%)'),
        ('3.0%', '3%', 0.03, 0.03, 'Band-order overturn — Band 6 > Band 5'),
        ('3.5%', '3%', 0.035, 0.03, 'Bounded — band order already overturned'),
        ('9.2%', '3%', 0.092, 0.03, 'TSAC component overturn for China'),
    ]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title = doc.add_heading("TSAC Break-Point Summary (SOSAC = 3%)", level=2)
    for run in title.runs:
        run.font.name = FONT
        run.font.size = Pt(11)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("142 eligible Parties | Fund size: USD 1,000 M | State/IPLC split: 50/50 | Spearman vs pure IUSAF")
    run.font.name = FONT
    run.font.size = Pt(8.5)
    run.font.italic = True

    doc.add_paragraph()

    headers = ['TSAC', 'SOSAC', 'IUSAF %', 'Spearman ρ', 'What Happens']
    n_cols = len(headers)
    n_rows = len(scenarios)
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.font.name = FONT
        run.font.size = Pt(8.5)
        run.font.bold = True
        set_cell_shading(cell, HEADER_BG)

    # Data
    for i, (tsac_label, sosac_label, beta, gamma, desc) in enumerate(scenarios):
        m = compute_scenario(con, beta, gamma)
        row = table.rows[i + 1]

        values = [tsac_label, sosac_label, f"{m['iusaf_pct']:.1f}%", f"{m['spearman']:.3f}", desc]

        for j, val in enumerate(values):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 4 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(8.5)

            # Highlight Gini-minimum row
            if '2.5%' in tsac_label and 'Gini' in desc:
                run.font.bold = True
                set_cell_shading(cell, "D4EDDA")
            # Highlight overturn row
            elif '3.0%' in tsac_label and 'overturn' in desc.lower():
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.font.bold = True

    # Column widths
    widths = [Cm(2.0), Cm(2.0), Cm(2.5), Cm(2.5), Cm(7.5)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w

    doc.save(path)
    print(f"Saved: {path}")


def write_band_order_csv(con, path):
    """Generate the band-order preservation CSV."""
    scenarios = [
        ('0% (Pure IUSAF)', 0.0, 0.0),
        ('1.5% (Strict)', 0.015, 0.03),
        ('2.5% (Gini-minimum)', 0.025, 0.03),
        ('3.0% (Band-order overturn)', 0.03, 0.03),
        ('3.5% (Bounded)', 0.035, 0.03),
        ('9.2% (TSAC component overturn)', 0.092, 0.03),
    ]
    rows = []
    for label, beta, gamma in scenarios:
        m = compute_scenario(con, beta, gamma)
        preserved = ("YES (margin {:.1f}%)".format(m['margin'])
                     if m['order_ok'] and m['margin'] < 10 else
                     "YES" if m['order_ok'] else
                     "NO — Band 6 overtakes Band 5")
        rows.append({
            'TSAC Level': label,
            'Band 6 mean China (USD M)': round(m['b6_mean'], 2),
            'Band 5 mean Brazil India Mexico (USD M)': round(m['b5_mean'], 2),
            'Band 5 vs Band 6 margin (%)': round(m['margin'], 1),
            'IUSAF Band Order Preserved': preserved,
        })
    pd.DataFrame(rows).to_csv(path, index=False)
    print(f"Saved: {path}")


def write_breakpoint_summary_csv(con, path):
    """Generate the break-point summary CSV."""
    scenarios = [
        ('0%', '3%', 0.0, 0.03, 'SOSAC only — modest rank shift among SIDS'),
        ('1.5%', '3%', 0.015, 0.03, 'Strict — IUSAF dominant for all Parties'),
        ('2.5%', '3%', 0.025, 0.03, 'Gini-minimum — band order preserved (margin 5.4%)'),
        ('3.0%', '3%', 0.03, 0.03, 'Band-order overturn — Band 6 > Band 5'),
        ('3.5%', '3%', 0.035, 0.03, 'Bounded — band order already overturned'),
        ('9.2%', '3%', 0.092, 0.03, 'TSAC component overturn for China'),
    ]
    rows = []
    for tsac_label, sosac_label, beta, gamma, desc in scenarios:
        m = compute_scenario(con, beta, gamma)
        rows.append({
            'TSAC': tsac_label,
            'SOSAC': sosac_label,
            'IUSAF %': round(m['iusaf_pct'], 1),
            'Spearman rho': round(m['spearman'], 3),
            'What Happens': desc,
        })
    pd.DataFrame(rows).to_csv(path, index=False)
    print(f"Saved: {path}")


def main():
    con = duckdb.connect(database=':memory:')
    load_data(con)

    os.makedirs(OUT_DIR, exist_ok=True)

    # Word tables
    write_band_order_table(con, os.path.join(OUT_DIR, "iusaf-band-order-preservation.docx"))
    write_breakpoint_summary_table(con, os.path.join(OUT_DIR, "iusaf-breakpoint-summary.docx"))

    # CSV companions
    write_band_order_csv(con, os.path.join(OUT_DIR, "iusaf-band-order-preservation.csv"))
    write_breakpoint_summary_csv(con, os.path.join(OUT_DIR, "iusaf-breakpoint-summary.csv"))


if __name__ == "__main__":
    import duckdb
    main()
