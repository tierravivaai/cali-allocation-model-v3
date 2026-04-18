"""Generate a clean black-and-white Word table from the IUSAF UN-region CSV."""

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

CSV_PATH = "/Users/pauloldham/Documents/cali-allocation-model-unscale-v3/model-tables/iusaf-unregion-15042026.csv"
OUT_PATH = "/Users/pauloldham/Documents/cali-allocation-model-unscale-v3/model-tables/iusaf-unregion-15042026.docx"

FONT = "Times New Roman"
FONT_SIZE = Pt(10)

# --- Read & clean data ---
df = pd.read_csv(CSV_PATH)

col_map = {
    "region": "UN Region",
    "Countries (number)": "Countries",
    "total_allocation": "Total Allocation\n(USD M)",
    "state_component": "State Component\n(USD M)",
    "iplc_component": "IPLC Component\n(USD M)",
}
df = df.rename(columns=col_map)

for c in ["Total Allocation\n(USD M)", "State Component\n(USD M)", "IPLC Component\n(USD M)"]:
    df[c] = df[c].round(2)

# --- Build document ---
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title
title = doc.add_heading("IUSAF Allocation by UN Region", level=2)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = FONT
    run.font.size = Pt(13)
    run.font.color.rgb = None  # black

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Fund size: USD 1,000 M  |  State/IPLC split: 50/50")
run.font.name = FONT
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()

# --- Create table ---
n_rows = len(df)
n_cols = len(df.columns)
table = doc.add_table(rows=n_rows + 1, cols=n_cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# --- Header row (grey background, black text) ---
HEADER_BG = "D9D9D9"
for j, col_name in enumerate(df.columns):
    cell = table.rows[0].cells[j]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = col_name.split("\n")
    for k, line in enumerate(lines):
        if k > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = FONT
        run.font.size = FONT_SIZE
        run.font.bold = True
    set_cell_shading(cell, HEADER_BG)

# --- Data rows (plain black on white) ---
for i, row in df.iterrows():
    is_total = str(row.iloc[0]).strip().lower() == "total"
    for j, val in enumerate(row):
        cell = table.rows[i + 1].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        if j == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        display = val
        if isinstance(val, float):
            display = f"{val:,.2f}"
        elif isinstance(val, (int,)):
            display = str(val)

        run = p.add_run(str(display))
        run.font.name = FONT
        run.font.size = FONT_SIZE
        if is_total:
            run.font.bold = True

# --- Column widths ---
widths = [Cm(3.0), Cm(2.2), Cm(3.2), Cm(3.2), Cm(3.2)]
for row in table.rows:
    for j, w in enumerate(widths):
        row.cells[j].width = w

doc.save(OUT_PATH)
print(f"Saved → {OUT_PATH}")
