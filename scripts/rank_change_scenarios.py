"""Generate a Word table showing how party rankings shift across break-point scenarios."""

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE = "/Users/pauloldham/Documents/cali-allocation-model-unscale-v3/band-analysis/break-points/scenario_results"
OUT = "/Users/pauloldham/Documents/cali-allocation-model-unscale-v3/model-tables/iusaf-rank-change-scenarios.docx"

FONT = "Times New Roman"
HEADER_BG = "D9D9D9"

# Scenario definitions (name, filename, beta, gamma)
SCENARIOS = [
    ("Pure IUSAF", "pure_iusaf.csv", 0.0, 0.0),
    ("Strict\n(β=1.5%)", "strict.csv", 0.015, 0.03),
    ("Gini-minimum\n(β=2.5%)", "gini_minimum.csv", 0.025, 0.03),
    ("Band-order\noverturn\n(β=3.0%)", "band_order_overturn.csv", 0.03, 0.03),
    ("Bounded\n(β=3.5%)", "bounded.csv", 0.035, 0.03),
    ("TSAC Overturn\n(β≈9.2%)", "tsac_overturn.csv", 0.092, 0.03),
    ("SOSAC Overturn\n(γ≈12.5%)", "sosac_overturn.csv", 0.0, 0.125),
]

def load_ranked(path):
    df = pd.read_csv(path)
    df = df[df["party"].str.strip().str.lower() != "total"].copy()
    df = df.sort_values("total_allocation", ascending=False).reset_index(drop=True)
    # Dense rank (ties get same rank)
    df["rank"] = df["total_allocation"].rank(method="dense", ascending=False).astype(int)
    return df[["party", "total_allocation", "rank", "un_band"]].copy()

# Load all scenarios
scenario_dfs = {}
for name, fname, _, _ in SCENARIOS:
    scenario_dfs[name.replace("\n", " ")] = load_ranked(f"{BASE}/{fname}")

# Identify top 20 parties in the baseline (Pure IUSAF)
baseline = scenario_dfs["Pure IUSAF"]
baseline_top20_cutoff = baseline.iloc[19]["total_allocation"]
top20_parties = baseline[baseline["total_allocation"] >= baseline_top20_cutoff]["party"].tolist()

# Also add a few key large countries that may rise dramatically
key_risers = ["China", "Brazil", "India", "Mexico", "Argentina", "Kazakhstan",
              "Indonesia", "Iran (Islamic Republic of)", "Democratic Republic of the Congo",
              "Algeria", "Sudan", "Colombia"]
for p in key_risers:
    if p not in top20_parties:
        top20_parties.append(p)

# Build wide-form rank table
rows = []
for party in top20_parties:
    row = {"Party": party}
    band = None
    for sname, sdf in scenario_dfs.items():
        match = sdf[sdf["party"] == party]
        if len(match) > 0:
            row[f"Rank\n{sname}"] = int(match.iloc[0]["rank"])
            row[f"Alloc\n{sname}"] = round(match.iloc[0]["total_allocation"], 2)
            if band is None:
                band = match.iloc[0]["un_band"]
        else:
            row[f"Rank\n{sname}"] = None
            row[f"Alloc\n{sname}"] = None
    row["UN Band"] = band if band else ""
    rows.append(row)

# Sort by baseline rank
rows.sort(key=lambda r: r.get("Rank\nPure IUSAF", 999) or 999)

# --- Build document ---
doc = Document()
for section in doc.sections:
    section.orientation = 1  # landscape
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# Title
title = doc.add_heading("IUSAF Rank Changes Across Break-Point Scenarios", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = FONT
    run.font.size = Pt(13)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(
    "Top 20 baseline recipients + key large-area countries  |  "
    "Ranks by total allocation (1 = highest recipient)\n"
    "Shaded cells mark rank improvement of ≥5 positions vs Pure IUSAF; "
    "bold red marks rank worsening of ≥5 positions"
)
run.font.name = FONT
run.font.size = Pt(8.5)
run.font.italic = True

doc.add_paragraph()

# --- Table with just ranks (more readable) ---
scenario_names = [n.replace("\n", " ") for n, _, _, _ in SCENARIOS]
rank_cols = [f"Rank\n{n}" for n in scenario_names]

# Use only rank columns + party + band
display_cols = ["Party", "UN Band"] + rank_cols
data_for_table = []
for r in rows:
    row_data = {}
    row_data["Party"] = r["Party"]
    row_data["UN Band"] = r["UN Band"]
    for rc in rank_cols:
        row_data[rc] = r[rc]
    data_for_table.append(row_data)

n_rows = len(data_for_table)
n_cols = len(display_cols)
table = doc.add_table(rows=n_rows + 1, cols=n_cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

# Header row
for j, col in enumerate(display_cols):
    cell = table.rows[0].cells[j]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = col.split("\n")
    for k, line in enumerate(lines):
        if k > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = FONT
        run.font.size = Pt(7)
        run.font.bold = True
    set_cell_shading(cell, HEADER_BG)

# Baseline ranks for comparison
baseline_ranks = {r["Party"]: r.get("Rank\nPure IUSAF") for r in rows}

# Data rows
for i, row_data in enumerate(data_for_table):
    party = row_data["Party"]
    base_rank = baseline_ranks.get(party)
    for j, col in enumerate(display_cols):
        cell = table.rows[i + 1].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        if j <= 1:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        val = row_data.get(col)
        if val is not None and col.startswith("Rank"):
            display = str(int(val))
        elif val is not None:
            display = str(val)
        else:
            display = "—"

        run = p.add_run(display)
        run.font.name = FONT
        run.font.size = Pt(7)

        # Highlight rank changes
        if col.startswith("Rank") and base_rank is not None and val is not None:
            delta = int(val) - int(base_rank)  # positive = worse rank (higher number)
            if delta <= -5:
                # Improved by 5+ positions
                set_cell_shading(cell, "D4EDDA")  # green tint
                run.font.bold = True
            elif delta >= 5:
                # Worsened by 5+ positions
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

# Column widths
party_w = Cm(4.5)
band_w = Cm(3.5)
rank_w = Cm(2.8)
widths = [party_w, band_w] + [rank_w] * len(rank_cols)
for row in table.rows:
    for j, w in enumerate(widths):
        row.cells[j].width = w

doc.save(OUT)
print(f"Saved → {OUT}")
print(f"Parties: {n_rows}, Scenarios: {len(scenario_names)}")
