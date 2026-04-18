"""Generate two Word documents with ranked panels for break-point scenarios.

Doc 1: Strict / Bounded / Even — three separate ranked tables
Doc 2: TSAC Overturn / SOSAC Overturn — two separate ranked tables

Each panel: Rank | Party | Total Allocation (USD M) | UN Band
Only non-zero allocation parties included.
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE = "/Users/pauloldham/Documents/cali-allocation-model-unscale-v3/band-analysis/break-points/scenario_results"
OUT1 = "/Users/pauloldham/Documents/cali-allocation-model-unscale-v3/model-tables/iusaf-scenarios-strict-bounded-even.docx"
OUT2 = "/Users/pauloldham/Documents/cali-allocation-model-unscale-v3/model-tables/iusaf-scenarios-overturn.docx"

FONT = "Times New Roman"
HEADER_BG = "D9D9D9"
ROW_HIGHLIGHT = "E8F5E9"  # light green for SIDS
FONT_SIZE = Pt(7.5)

DOC1_SCENARIOS = [
    ("Strict (β=1.5%, γ=3%)", "strict.csv"),
    ("Bounded (β=3.5%, γ=3%)", "bounded.csv"),
    ("Even (β=5.0%, γ=3%)", "even.csv"),
]

DOC2_SCENARIOS = [
    ("TSAC Overturn (β≈9.2%, γ=3%)", "tsac_overturn.csv"),
    ("SOSAC Overturn (β=0%, γ≈12.5%)", "sosac_overturn.csv"),
]


def load_ranked(path):
    df = pd.read_csv(path)
    df = df[df["party"].str.strip().str.lower() != "total"].copy()
    df = df[df["total_allocation"] > 0].copy()
    df = df.sort_values("total_allocation", ascending=False).reset_index(drop=True)
    df["rank"] = df["total_allocation"].rank(method="dense", ascending=False).astype(int)
    return df[["party", "total_allocation", "rank", "un_band", "is_sids"]].copy()


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_panel(doc, data, title, subtitle=None):
    """Add one ranked table panel to the document."""
    h = doc.add_heading(title, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = FONT
        run.font.size = Pt(11)

    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(subtitle)
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.italic = True

    n_rows = len(data)
    n_cols = 4
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Rank", "Party", "Total Allocation\n(USD M)", "UN Band"]
    for j, hdr in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = hdr.split("\n")
        for k, line in enumerate(lines):
            if k > 0:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.name = FONT
            run.font.size = Pt(7)
            run.font.bold = True
        set_cell_shading(cell, HEADER_BG)

    for i, (_, row) in enumerate(data.iterrows()):
        is_sids = row.get("is_sids", False)

        # Rank
        cell = table.rows[i + 1].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(int(row["rank"])))
        run.font.name = FONT
        run.font.size = FONT_SIZE

        # Party
        cell = table.rows[i + 1].cells[1]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(row["party"]))
        run.font.name = FONT
        run.font.size = FONT_SIZE
        if is_sids:
            run.font.bold = True

        # Allocation
        cell = table.rows[i + 1].cells[2]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{row['total_allocation']:,.2f}")
        run.font.name = FONT
        run.font.size = FONT_SIZE

        # Band
        cell = table.rows[i + 1].cells[3]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(row["un_band"]))
        run.font.name = FONT
        run.font.size = FONT_SIZE

        # Highlight SIDS rows
        if is_sids:
            for j in range(n_cols):
                set_cell_shading(table.rows[i + 1].cells[j], ROW_HIGHLIGHT)

    # Column widths
    widths = [Cm(1.2), Cm(5.5), Cm(3.0), Cm(4.5)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w

    doc.add_paragraph()  # spacer


def build_doc(out_path, scenarios, doc_title, doc_subtitle):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title = doc.add_heading(doc_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = FONT
        run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(doc_subtitle)
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.italic = True

    # Legend
    legend = doc.add_paragraph()
    legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = legend.add_run("Bold name + green row = SIDS  |  Ranked by total allocation descending  |  Zero-allocation parties excluded")
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.italic = True

    doc.add_paragraph()

    for name, fname in scenarios:
        df = load_ranked(f"{BASE}/{fname}")
        add_panel(doc, df, name, f"{len(df)} parties with non-zero allocation")

    doc.save(out_path)
    print(f"Saved → {out_path}")


build_doc(
    OUT1, DOC1_SCENARIOS,
    "IUSAF Allocation: Negotiation Scenarios",
    "Fund size: USD 1,000 M  |  State/IPLC split: 50/50  |  SIDS highlighted"
)

build_doc(
    OUT2, DOC2_SCENARIOS,
    "IUSAF Allocation: Overturn Scenarios",
    "Fund size: USD 1,000 M  |  State/IPLC split: 50/50  |  Beyond negotiation space"
)
